import io
import json
import time
from unittest import mock

import pytest
from fastapi.testclient import TestClient

from app import main, storage


@pytest.fixture
def client(tmp_path, monkeypatch):
    monkeypatch.setattr(storage, "DATA_DIR", tmp_path / "data")
    storage.DATA_DIR.mkdir(parents=True, exist_ok=True)

    async def fake_run(meeting_id, cfg, task_id=None):
        storage.update_meta(meeting_id, status="done")
    monkeypatch.setattr(main.tasks, "run_pipeline", fake_run)

    with TestClient(main.app) as c:
        yield c


def test_index_html(client):
    r = client.get("/")
    assert r.status_code == 200
    assert "text/html" in r.headers["content-type"]


def test_upload_and_list(client):
    audio_bytes = b"fake-audio"
    files = {"audio": ("m.m4a", io.BytesIO(audio_bytes), "audio/m4a")}
    data = {"title": "Demo Meeting"}
    r = client.post("/api/upload", files=files, data=data)
    assert r.status_code == 200
    body = r.json()
    assert "meeting_id" in body and "task_id" in body
    meetings = client.get("/api/meetings").json()
    assert any(m["id"] == body["meeting_id"] for m in meetings)


def test_upload_rejects_non_audio(client):
    files = {"audio": ("t.txt", io.BytesIO(b"x"), "text/plain")}
    data = {"title": "x"}
    r = client.post("/api/upload", files=files, data=data)
    assert r.status_code == 415


def test_upload_streams_to_disk_and_enforces_size_limit(client, monkeypatch, tmp_path):
    # S3/P2: upload now streams to disk in chunks (no full-file read into memory).
    # Sanity: a normal small upload still works.
    files = {"audio": ("a.wav", io.BytesIO(b"\x00" * 100), "audio/wav")}
    r = client.post("/api/upload", files=files, data={"title": "T"})
    assert r.status_code == 200

    # Force a tiny limit: exceeding it must 413 (not OOM), and leave no tmp file.
    monkeypatch.setattr(main, "MAX_UPLOAD_BYTES", 8)
    files = {"audio": ("b.wav", io.BytesIO(b"\x00" * 100), "audio/wav")}
    r = client.post("/api/upload", files=files, data={"title": "T2"})
    assert r.status_code == 413
    uploads_dir = storage.DATA_DIR / "uploads"
    # the aborted tmp upload must have been cleaned up
    leftover = [p for p in uploads_dir.glob("upload_*") if p.is_file()]
    assert leftover == []


def test_get_meeting(client):
    files = {"audio": ("a.wav", io.BytesIO(b"x"), "audio/wav")}
    r = client.post("/api/upload", files=files, data={"title": "T"})
    mid = r.json()["meeting_id"]
    detail = client.get(f"/api/meetings/{mid}").json()
    assert detail["meta"]["id"] == mid


def test_delete_meeting(client):
    files = {"audio": ("a.wav", io.BytesIO(b"x"), "audio/wav")}
    mid = client.post("/api/upload", files=files, data={"title": "T"}).json()["meeting_id"]
    r = client.delete(f"/api/meetings/{mid}")
    assert r.status_code == 200
    assert client.get(f"/api/meetings/{mid}").status_code == 404


def test_delete_meeting_keep_moves_to_trash(client):
    mid = client.post("/api/upload", files={"audio": ("a.wav", io.BytesIO(b"x"), "audio/wav")}, data={"title": "T"}).json()["meeting_id"]
    r = client.delete(f"/api/meetings/{mid}?keep=1")
    assert r.status_code == 200
    body = r.json()
    assert body["trashed"] is True
    # gone from history list
    assert not any(m["id"] == mid for m in client.get("/api/meetings").json())
    # files preserved under visible 回收站 dir
    trash = storage.trash_dir()
    assert trash.name == "回收站"
    moved = list(trash.iterdir())
    assert len(moved) == 1
    assert (moved[0] / "meta.json").exists()
    assert body["trash_path"].endswith("回收站/" + moved[0].name)


def test_delete_meeting_full_removes_files(client):
    mid = client.post("/api/upload", files={"audio": ("a.wav", io.BytesIO(b"x"), "audio/wav")}, data={"title": "T"}).json()["meeting_id"]
    r = client.delete(f"/api/meetings/{mid}")
    assert r.json()["trashed"] is False
    assert not storage.meeting_dir(mid).exists()


def test_rename_meeting(client):
    mid = client.post("/api/upload", files={"audio": ("a.wav", io.BytesIO(b"x"), "audio/wav")}, data={"title": "旧名"}).json()["meeting_id"]
    r = client.put(f"/api/meetings/{mid}/title", json={"title": "新名字"})
    assert r.status_code == 200 and r.json()["title"] == "新名字"
    assert client.get(f"/api/meetings/{mid}").json()["meta"]["title"] == "新名字"


def test_rename_meeting_rejects_empty(client):
    mid = client.post("/api/upload", files={"audio": ("a.wav", io.BytesIO(b"x"), "audio/wav")}, data={"title": "T"}).json()["meeting_id"]
    assert client.put(f"/api/meetings/{mid}/title", json={"title": "  "}).status_code == 400


def test_set_tags(client):
    mid = client.post("/api/upload", files={"audio": ("a.wav", io.BytesIO(b"x"), "audio/wav")}, data={"title": "T"}).json()["meeting_id"]
    r = client.put(f"/api/meetings/{mid}/tags", json={"tags": ["周会", "产品", "周会"]})
    assert r.status_code == 200
    assert r.json()["tags"] == ["周会", "产品"]  # de-duped, order kept
    meta = client.get(f"/api/meetings/{mid}").json()["meta"]
    assert meta["tags"] == ["周会", "产品"]


def test_set_tags_rejects_bad_shape(client):
    mid = client.post("/api/upload", files={"audio": ("a.wav", io.BytesIO(b"x"), "audio/wav")}, data={"title": "T"}).json()["meeting_id"]
    assert client.put(f"/api/meetings/{mid}/tags", json={"tags": "not-a-list"}).status_code == 400


def test_trash_not_listed_as_meeting(client):
    # 回收站 dir itself must not appear in the meeting list
    storage.trash_dir().mkdir(parents=True, exist_ok=True)
    meetings = client.get("/api/meetings").json()
    assert all("回收站" not in m["id"] for m in meetings)


def test_rename_speakers_persists(client):
    mid = client.post("/api/upload", files={"audio": ("a.wav", io.BytesIO(b"x"), "audio/wav")}, data={"title": "T"}).json()["meeting_id"]
    r = client.put(f"/api/meetings/{mid}/speakers", json={"names": {"0": "Alice", "1": "Bob"}})
    assert r.status_code == 200
    meta = client.get(f"/api/meetings/{mid}").json()["meta"]
    assert meta["speaker_names"] == {"0": "Alice", "1": "Bob"}


# --- P0.3: speaker merge / split (pure meta remap) ---
def _seed_speakers(client, mid):
    storage.save_raw(mid, {"sentences": [
        {"start": 0, "end": 1000, "spk": 0, "text": "a1"},
        {"start": 1100, "end": 2000, "spk": 1, "text": "b1"},
        {"start": 2100, "end": 3000, "spk": 2, "text": "c1"},
        {"start": 3100, "end": 4000, "spk": 1, "text": "b2"},
    ], "spk_count": 3})
    client.put(f"/api/meetings/{mid}/speakers", json={"names": {"0": "Alice", "1": "Bob", "2": "Carol"}})


def test_remap_merge_speakers(client):
    mid = _upload(client)
    _seed_speakers(client, mid)
    # merge spk 2 -> 0 (Carol into Alice)
    r = client.post(f"/api/meetings/{mid}/speakers/remap", json={"map": {"2": "0"}})
    assert r.status_code == 200
    raw = storage.get_meeting(mid)["raw"]["sentences"]
    # distinct spks now {0,1} -> normalized {0,1}; spk 2 sentences became 0
    assert sorted({s["spk"] for s in raw}) == [0, 1]
    assert raw[0]["spk"] == 0 and raw[2]["spk"] == 0  # Carol's sentence now Alice
    meta = storage.get_meeting(mid)["meta"]
    assert meta["spk_count"] == 2
    # names carried over; spk 0 keeps Alice
    assert meta["speaker_names"]["0"] == "Alice"


def test_remap_sentence_reassign_then_renormalize(client):
    mid = _upload(client)
    _seed_speakers(client, mid)
    # spk 1 owns two sentences (idx 1 & 3); split by moving idx 1 to a brand-new
    # speaker id. distinct grows 3 -> 4 and ids renormalize to contiguous 0..3.
    r = client.post(f"/api/meetings/{mid}/speakers/remap",
                    json={"sentences": {"1": "9"}})
    assert r.status_code == 200
    raw = storage.get_meeting(mid)["raw"]["sentences"]
    assert sorted({s["spk"] for s in raw}) == [0, 1, 2, 3]  # 4 distinct, contiguous
    meta = storage.get_meeting(mid)["meta"]
    assert meta["spk_count"] == 4
    # the moved sentence and the remaining spk-1 sentence are now different speakers
    assert raw[1]["spk"] != raw[3]["spk"]


def test_remap_updates_processed_md_speaker_refs(client):
    mid = _upload(client)
    _seed_speakers(client, mid)
    storage.save_processed(mid, "## 说话人 0\n\na\n\n## 说话人 2\n\nc\n")
    client.post(f"/api/meetings/{mid}/speakers/remap", json={"map": {"2": "0"}})
    processed = storage.get_meeting(mid)["processed"]
    # 说话人 2 merged into 0; the header reference should be remapped
    assert "说话人 2" not in processed
    assert processed.count("说话人 0") == 2


def test_remap_rejects_bad_sentence_index(client):
    mid = _upload(client)
    _seed_speakers(client, mid)
    assert client.post(f"/api/meetings/{mid}/speakers/remap",
                       json={"sentences": {"99": "0"}}).status_code == 400


def test_set_meeting_context_persists(client):
    mid = client.post("/api/upload", files={"audio": ("a.wav", io.BytesIO(b"x"), "audio/wav")}, data={"title": "T"}).json()["meeting_id"]
    r = client.put(f"/api/meetings/{mid}/context", json={"meeting_context": "X 项目周会；术语：K8s"})
    assert r.status_code == 200
    meta = client.get(f"/api/meetings/{mid}").json()["meta"]
    assert meta["meeting_context"] == "X 项目周会；术语：K8s"


def test_templates_seed_and_roundtrip(client):
    d = client.get("/api/settings/templates").json()
    tpls = d["templates"]
    assert any(t["id"] == "general" for t in tpls)
    assert any(t["name"] == "周会" for t in tpls)
    client.put("/api/settings/templates", json={"templates": tpls + [{"name": "播客剪辑", "direction": "话题脉络"}]})
    d2 = client.get("/api/settings/templates").json()
    assert any(t["name"] == "播客剪辑" for t in d2["templates"])
    assert d2["templates"][-1]["id"].startswith("c-")
    assert d2["templates"][-1]["direction"] == "话题脉络"


def test_set_meeting_template(client):
    mid = client.post("/api/upload", files={"audio": ("a.wav", io.BytesIO(b"x"), "audio/wav")}, data={"title": "T"}).json()["meeting_id"]
    r = client.put(f"/api/meetings/{mid}/context", json={"template": "company"})
    assert r.status_code == 200
    assert client.get(f"/api/meetings/{mid}").json()["meta"]["template"] == "company"


def test_onboard_flag_roundtrip(client, monkeypatch):
    monkeypatch.setattr(main, "_persist_storage", lambda **kw: None)
    prev = main.config.storage.onboarded
    try:
        main.config.storage.onboarded = False
        assert client.get("/api/settings").json()["onboarded"] is False
        assert client.post("/api/settings/onboard").status_code == 200
        assert client.get("/api/settings").json()["onboarded"] is True
    finally:
        main.config.storage.onboarded = prev


def test_save_data_dir_marks_onboarded(client, monkeypatch):
    monkeypatch.setattr(main, "_persist_storage", lambda **kw: None)
    prev = main.config.storage.onboarded
    try:
        main.config.storage.onboarded = False
        r = client.post("/api/settings", json={"data_dir": str(storage.DATA_DIR)})
        assert r.status_code == 200
        assert client.get("/api/settings").json()["onboarded"] is True
    finally:
        main.config.storage.onboarded = prev


def test_export_txt_uses_speaker_names(client):
    mid = client.post("/api/upload", files={"audio": ("a.wav", io.BytesIO(b"x"), "audio/wav")}, data={"title": "T"}).json()["meeting_id"]
    storage.save_raw(mid, {"sentences": [
        {"start": 0, "end": 1000, "spk": 0, "text": "hi"},
        {"start": 1100, "end": 2000, "spk": 1, "text": "yo"},
    ]})
    client.put(f"/api/meetings/{mid}/speakers", json={"names": {"0": "Alice", "1": "Bob"}})
    txt = client.get(f"/api/meetings/{mid}/export?format=txt").text
    assert "Alice" in txt and "Bob" in txt
    assert "说话人0" not in txt and "说话人1" not in txt


def test_export_srt_uses_speaker_names(client):
    mid = client.post("/api/upload", files={"audio": ("a.wav", io.BytesIO(b"x"), "audio/wav")}, data={"title": "T"}).json()["meeting_id"]
    storage.save_raw(mid, {"sentences": [{"start": 0, "end": 1000, "spk": 0, "text": "hi"}]})
    client.put(f"/api/meetings/{mid}/speakers", json={"names": {"0": "Alice"}})
    srt = client.get(f"/api/meetings/{mid}/export?format=srt").text
    assert "[Alice]" in srt


def test_export_md_uses_speaker_names(client):
    mid = client.post("/api/upload", files={"audio": ("a.wav", io.BytesIO(b"x"), "audio/wav")}, data={"title": "T"}).json()["meeting_id"]
    storage.save_processed(mid, "## 说话人 0\n\n你好\n\n## 说话人1\n\n嗯\n")
    client.put(f"/api/meetings/{mid}/speakers", json={"names": {"0": "Alice", "1": "Bob"}})
    md = client.get(f"/api/meetings/{mid}/export?format=md").text
    assert "Alice" in md and "Bob" in md
    assert "说话人0" not in md and "说话人1" not in md


def test_export_md_before_polish_returns_400(client):
    mid = client.post("/api/upload", files={"audio": ("a.wav", io.BytesIO(b"x"), "audio/wav")}, data={"title": "T"}).json()["meeting_id"]
    r = client.get(f"/api/meetings/{mid}/export?format=md")
    assert r.status_code == 400


# --- P1.1 / P1.2: docx, minutes, export options ---
def test_export_docx_renders_processed(client):
    from docx import Document
    mid = _upload(client)
    storage.save_processed(mid, "## 说话人 0\n\n你好世界\n\n- 要点一\n")
    client.put(f"/api/meetings/{mid}/speakers", json={"names": {"0": "Alice"}})
    r = client.get(f"/api/meetings/{mid}/export?format=docx")
    assert r.status_code == 200
    assert r.headers["content-type"].startswith(_docx_prefix())
    doc = Document(io.BytesIO(r.content))
    paras = [p.text for p in doc.paragraphs]
    assert any("Alice" in t for t in paras)      # speaker name applied
    assert any("要点一" in t for t in paras)     # bullet preserved


def test_export_minutes_uses_structured_summary(client):
    mid = _upload(client)
    storage.save_summary_json(mid, {"summary": "概述X", "decisions": ["决定甲"],
                                    "action_items": ["待办乙"], "open_questions": []})
    client.put(f"/api/meetings/{mid}/speakers", json={"names": {"0": "Alice"}})
    md = client.get(f"/api/meetings/{mid}/export?format=minutes").text
    assert "概述X" in md and "决定甲" in md and "待办乙" in md
    assert md.startswith("# ")  # title heading prepended


def test_export_minutes_docx(client):
    from docx import Document
    mid = _upload(client)
    storage.save_summary_json(mid, {"summary": "概述X", "decisions": [], "action_items": [], "open_questions": []})
    r = client.get(f"/api/meetings/{mid}/export?format=minutes_docx")
    assert r.status_code == 200
    doc = Document(io.BytesIO(r.content))
    assert any("概述X" in p.text for p in doc.paragraphs)


def test_export_txt_options_omit_speakers_and_timestamps(client):
    mid = _upload(client)
    storage.save_raw(mid, {"sentences": [
        {"start": 0, "end": 1000, "spk": 0, "text": "你好"},
        {"start": 1100, "end": 2000, "spk": 1, "text": "再见"},
    ]})
    client.put(f"/api/meetings/{mid}/speakers", json={"names": {"0": "Alice", "1": "Bob"}})
    full = client.get(f"/api/meetings/{mid}/export?format=txt").text
    assert "Alice" in full and "00:00:00" in full
    no_spk = client.get(f"/api/meetings/{mid}/export?format=txt&speakers=false").text
    assert "Alice" not in no_spk and "你好" in no_spk
    no_ts = client.get(f"/api/meetings/{mid}/export?format=txt&timestamps=false").text
    assert "00:00:00" not in no_ts and "Alice" in no_ts
    bare = client.get(f"/api/meetings/{mid}/export?format=txt&speakers=false&timestamps=false").text
    assert bare.strip() == "你好\n再见"


def test_export_rejects_unknown_format(client):
    mid = _upload(client)
    assert client.get(f"/api/meetings/{mid}/export?format=pdf").status_code == 400


def _docx_prefix():
    return "application/vnd.openxmlformats-officedocument"


def _upload(client, title="T"):
    return client.post("/api/upload", files={"audio": ("a.wav", io.BytesIO(b"x"), "audio/wav")}, data={"title": title}).json()["meeting_id"]


def test_retry_llm_runs_in_background_with_registered_task(client, monkeypatch):
    calls = {}

    async def fake_retry(meeting_id, cfg, task_id=None):
        calls["task_id"] = task_id
        storage.update_meta(meeting_id, status="done")

    monkeypatch.setattr(main.tasks, "retry_llm", fake_retry)
    mid = _upload(client)
    r = client.post(f"/api/meetings/{mid}/retry-llm")
    assert r.status_code == 200
    tid = r.json()["task_id"]
    # the task_id handed to the frontend is the one the pipeline updates
    assert calls["task_id"] == tid
    assert client.get(f"/api/tasks/{tid}").status_code == 200


def test_resume_llm_stage_uses_background_task(client, monkeypatch):
    async def fake_retry(meeting_id, cfg, task_id=None):
        storage.update_meta(meeting_id, status="done")

    monkeypatch.setattr(main.tasks, "retry_llm", fake_retry)
    mid = _upload(client)
    main.tasks._tasks.clear()  # simulate process restart (resume's use case)
    storage.update_meta(mid, status="llm_polishing")
    body = client.post(f"/api/meetings/{mid}/resume").json()
    assert body["ok"] is True
    assert body["action"] == "retry_llm"
    assert body["task_id"]
    assert client.get(f"/api/tasks/{body['task_id']}").status_code == 200


def test_resume_pipeline_stage_passes_registered_task(client, monkeypatch):
    seen = {}

    async def fake_run(meeting_id, cfg, task_id=None):
        seen["task_id"] = task_id
        storage.update_meta(meeting_id, status="done")

    monkeypatch.setattr(main.tasks, "run_pipeline", fake_run)
    mid = _upload(client)
    main.tasks._tasks.clear()  # simulate process restart (resume's use case)
    storage.update_meta(mid, status="asr_running")  # simulate crash mid-ASR
    body = client.post(f"/api/meetings/{mid}/resume").json()
    assert body["action"] == "run_pipeline"
    assert seen["task_id"] == body["task_id"]


def test_asr_settings_rejects_sidecar_until_v06(client, monkeypatch):
    monkeypatch.setattr(main, "_persist_asr_config", lambda: None)
    r = client.post("/api/settings/asr", json={"stream_engine": "sidecar"})
    assert r.status_code == 400
    r = client.post("/api/settings/asr", json={"stream_engine": "funasr"})
    assert r.status_code == 200


def test_trash_endpoints(client):
    mid = _upload(client)
    client.delete(f"/api/meetings/{mid}?keep=1")
    items = client.get("/api/trash").json()["items"]
    assert any(i["name"] == mid for i in items)
    assert client.post(f"/api/trash/{mid}/restore").status_code == 200
    assert client.get(f"/api/meetings/{mid}").status_code == 200
    # restore again -> 400 (not in trash anymore)
    assert client.post(f"/api/trash/{mid}/restore").status_code == 400
    # trash then permanently delete via trash endpoint
    client.delete(f"/api/meetings/{mid}?keep=1")
    assert client.delete(f"/api/trash/{mid}").status_code == 200
    assert client.get("/api/trash").json()["items"] == []
    assert client.delete(f"/api/trash/{mid}").status_code == 404


def test_edit_processed(client):
    mid = _upload(client)
    r = client.put(f"/api/meetings/{mid}/processed", json={"text": "# 整理后"})
    assert r.status_code == 200
    assert storage.get_meeting(mid)["processed"] == "# 整理后"
    assert client.put(f"/api/meetings/{mid}/processed", json={"text": "  "}).status_code == 400


def test_edit_summary_structured_and_markdown(client):
    mid = _upload(client)
    sj = {"summary": "概述", "decisions": ["d1"], "action_items": ["a1"], "open_questions": []}
    r = client.put(f"/api/meetings/{mid}/summary", json={"summary_json": sj})
    assert r.status_code == 200 and r.json()["structured"] is True
    data = storage.get_meeting(mid)
    assert data["summary_json"]["decisions"] == ["d1"]
    assert "概述" in data["summary"] and "d1" in data["summary"]
    # manual markdown edit drops the structured form
    r = client.put(f"/api/meetings/{mid}/summary", json={"text": "## 手动总结"})
    assert r.status_code == 200 and r.json()["structured"] is False
    data = storage.get_meeting(mid)
    assert data["summary_json"] is None
    assert data["summary"] == "## 手动总结"
    assert client.put(f"/api/meetings/{mid}/summary", json={"text": ""}).status_code == 400


def test_resume_live_recording_marks_error(client):
    mid = _upload(client)
    main.tasks._tasks.clear()  # simulate process restart mid-recording
    storage.update_meta(mid, status="live_recording")
    body = client.post(f"/api/meetings/{mid}/resume").json()
    assert body["ok"] is True and body["action"] == "mark_error"
    assert storage.get_meeting(mid)["meta"]["status"] == "error"


def test_browse_lists_only_subdirs(client, tmp_path):
    (tmp_path / "subA").mkdir()
    (tmp_path / "subB").mkdir()
    (tmp_path / "file.txt").write_text("x")
    body = client.get("/api/browse", params={"path": str(tmp_path)}).json()
    names = [d["name"] for d in body["dirs"]]
    assert "subA" in names and "subB" in names
    assert "file.txt" not in names
    assert body["writable"] is True
    assert body["parent"] is not None


def test_migrate_moves_meetings(client, tmp_path):
    src = tmp_path / "old_data"
    m = src / "m1"
    m.mkdir(parents=True)
    (m / "meta.json").write_text(json.dumps({"id": "m1", "title": "t"}))
    (m / "audio.wav").write_bytes(b"audio")
    (m / "sub").mkdir()
    (m / "sub" / "nested.json").write_text("{}")
    body = client.post("/api/settings/migrate", json={"from_dir": str(src)}).json()
    assert body["count"] == 1
    assert "m1" in body["moved"]
    dst_meeting = storage.get_data_dir() / "m1"
    assert (dst_meeting / "meta.json").exists()
    assert (dst_meeting / "audio.wav").exists()
    assert (dst_meeting / "sub" / "nested.json").exists()  # nested dirs copied
    assert not m.exists()  # source removed


def test_edit_sentence(client):
    files = {"audio": ("a.wav", io.BytesIO(b"x"), "audio/wav")}
    mid = client.post("/api/upload", files=files, data={"title": "T"}).json()["meeting_id"]
    storage.save_raw(mid, {"sentences": [{"start": 0, "end": 100, "spk": 0, "text": "旧文本"}]})
    r = client.put(f"/api/meetings/{mid}/sentence", json={"index": 0, "text": "新文本"})
    assert r.status_code == 200
    after = client.get(f"/api/meetings/{mid}").json()["raw"]["sentences"][0]["text"]
    assert after == "新文本"


def test_edit_sentence_rejects_bad_index(client):
    files = {"audio": ("a.wav", io.BytesIO(b"x"), "audio/wav")}
    mid = client.post("/api/upload", files=files, data={"title": "T"}).json()["meeting_id"]
    storage.save_raw(mid, {"sentences": [{"start": 0, "end": 100, "spk": 0, "text": "x"}]})
    assert client.put(f"/api/meetings/{mid}/sentence", json={"index": 5, "text": "y"}).status_code == 400
    assert client.put(f"/api/meetings/{mid}/sentence", json={"index": 0, "text": ""}).status_code == 400


def test_hotwords_roundtrip(client):
    r = client.post("/api/settings/hotwords", json={"hotwords": ["丁老师", "云南"]})
    assert r.status_code == 200 and r.json()["count"] == 2
    got = client.get("/api/settings/hotwords").json()["hotwords"]
    assert got == ["丁老师", "云南"]


def test_hotwords_dedup(client):
    r = client.post("/api/settings/hotwords", json={"hotwords": ["a", "b", "a", "b", "c"]})
    assert r.json()["count"] == 3 and r.json()["duplicates"] == 2
    got = client.get("/api/settings/hotwords").json()["hotwords"]
    assert got == ["a", "b", "c"]


# --- ASR status / unload ---
def test_asr_status(client):
    d = client.get("/api/asr/status").json()
    assert "loaded" in d and "rss_mb" in d and "idle_unload_minutes" in d


def test_asr_unload_when_idle(client, monkeypatch):
    main.tasks._tasks.clear()
    monkeypatch.setattr(main.asr, "_model", object())
    monkeypatch.setattr(main.asr, "_busy", False)
    r = client.post("/api/asr/unload")
    assert r.status_code == 200
    assert r.json()["unloaded"] is True


def test_asr_unload_refused_when_busy(client, monkeypatch):
    monkeypatch.setattr(main.asr, "_model", object())
    monkeypatch.setattr(main.asr, "_busy", True)
    assert client.post("/api/asr/unload").status_code == 409


def test_asr_unload_refused_when_task_pending(client, monkeypatch):
    monkeypatch.setattr(main.asr, "_model", object())
    monkeypatch.setattr(main.asr, "_busy", False)
    main.tasks._tasks["t1"] = {"status": "asr_running"}
    try:
        assert client.post("/api/asr/unload").status_code == 409
    finally:
        main.tasks._tasks.clear()


# --- port check / server settings ---
def test_port_check_free(client):
    import socket as _s
    sk = _s.socket(_s.AF_INET, _s.SOCK_STREAM)
    sk.bind(("0.0.0.0", 0))
    free_port = sk.getsockname()[1]
    sk.close()
    d = client.post("/api/settings/server/check", json={"port": free_port, "host": "0.0.0.0"}).json()
    assert d["ok"] is True and d["self"] is False


def test_port_check_conflict(client):
    import socket as _s
    sk = _s.socket(_s.AF_INET, _s.SOCK_STREAM)
    sk.setsockopt(_s.SOL_SOCKET, _s.SO_REUSEADDR, 0)
    sk.bind(("0.0.0.0", 0))
    sk.listen(1)
    held_port = sk.getsockname()[1]
    try:
        d = client.post("/api/settings/server/check", json={"port": held_port, "host": "0.0.0.0"}).json()
        assert d["ok"] is False
    finally:
        sk.close()


def test_port_check_self(client):
    d = client.post("/api/settings/server/check", json={"port": main._RUNNING_PORT, "host": "0.0.0.0"}).json()
    assert d["ok"] is True and d["self"] is True


def test_port_check_rejects_invalid(client):
    assert client.post("/api/settings/server/check", json={"port": 0}).status_code == 400
    assert client.post("/api/settings/server/check", json={"port": 70000}).status_code == 400


def test_server_settings_roundtrip(client, monkeypatch):
    monkeypatch.setattr(main, "_persist_server_config", lambda *a, **kw: None)
    prev = (main.config.server.host, main.config.server.port, main.config.asr.idle_unload_minutes)
    try:
        r = client.post("/api/settings/server", json={
            "port": 9999, "host": "127.0.0.1", "idle_unload_minutes": 5})
        assert r.status_code == 200
        d = r.json()
        assert d["ok"] and d["restart_required"] is True  # 9999 != running port
        assert main.config.server.port == 9999
        assert main.config.asr.idle_unload_minutes == 5
        got = client.get("/api/settings/server").json()
        assert got["port"] == 9999 and got["running_port"] == main._RUNNING_PORT
    finally:
        main.config.server.host, main.config.server.port = prev[0], prev[1]
        main.config.asr.idle_unload_minutes = prev[2]


# --- idle watcher decision logic ---
def test_idle_check_unloads_when_idle(client, monkeypatch):
    main.tasks._tasks.clear()
    monkeypatch.setattr(main.asr, "_model", object())
    monkeypatch.setattr(main.asr, "_busy", False)
    monkeypatch.setattr(main.asr, "_last_used", time.time() - 1000)
    assert main._idle_check(60) is True
    assert not main.asr.is_loaded()


def test_idle_check_skips_when_below_threshold(client, monkeypatch):
    main.tasks._tasks.clear()
    monkeypatch.setattr(main.asr, "_model", object())
    monkeypatch.setattr(main.asr, "_busy", False)
    monkeypatch.setattr(main.asr, "_last_used", time.time())
    assert main._idle_check(3600) is False
    assert main.asr.is_loaded()


def test_idle_check_skips_when_task_active(client, monkeypatch):
    monkeypatch.setattr(main.asr, "_model", object())
    monkeypatch.setattr(main.asr, "_busy", False)
    monkeypatch.setattr(main.asr, "_last_used", time.time() - 1000)
    main.tasks._tasks["t"] = {"status": "asr_running"}
    try:
        assert main._idle_check(60) is False
        assert main.asr.is_loaded()
    finally:
        main.tasks._tasks.clear()


def test_idle_watcher_loop_unloads(client, monkeypatch):
    """Full wiring: watcher task → _idle_check → unload_model, in a real loop."""
    import asyncio as _a
    main.tasks._tasks.clear()
    monkeypatch.setattr(main.asr, "_model", object())
    monkeypatch.setattr(main.asr, "_busy", False)
    monkeypatch.setattr(main.asr, "_last_used", time.time() - 1000)
    monkeypatch.setattr(main, "_WATCHER_TICK_S", 0.05)
    monkeypatch.setattr(main, "_idle_unload_seconds", lambda: 60)

    async def go():
        task = _a.create_task(main._idle_watcher())
        for _ in range(20):  # wait up to ~1s for the watcher to fire
            if not main.asr.is_loaded():
                task.cancel()
                break
            await _a.sleep(0.05)
        else:
            task.cancel()
        try:
            await task
        except _a.CancelledError:
            pass
    _a.run(go())
    assert not main.asr.is_loaded(), "watcher should have unloaded the model"
